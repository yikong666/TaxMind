from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from rag.parsing.chunker import ParentChildChunker
from rag.parsing.document_parser import DocumentParser, ParsedSection
from tests.test_knowledge_bases import authenticate, create_knowledge_base


def upload_document(
    client: TestClient,
    headers: dict[str, str],
    knowledge_base_id: int,
    filename: str,
    content: bytes,
    content_type: str,
) -> int:
    response = client.post(
        f"/api/v1/knowledge-bases/{knowledge_base_id}/documents",
        headers=headers,
        files={"files": (filename, content, content_type)},
    )
    assert response.status_code == 201
    return response.json()["data"][0]["id"]


def test_markdown_parser_preserves_heading_sections() -> None:
    sections = DocumentParser().parse(
        "policy.md", "# 总则\n第一条内容。\n## 优惠条件\n第二条内容。".encode()
    )
    assert [item.heading for item in sections] == ["总则", "优惠条件"]
    assert sections[1].content == "第二条内容。"


def test_docx_parser_extracts_heading_and_body() -> None:
    document = DocxDocument()
    document.add_heading("增值税政策", level=1)
    document.add_paragraph("适用于小规模纳税人。")
    stream = BytesIO()
    document.save(stream)

    sections = DocumentParser().parse("policy.docx", stream.getvalue())
    assert sections == [ParsedSection("增值税政策", "适用于小规模纳税人。")]


def test_parent_child_chunker_respects_sizes_and_mapping() -> None:
    content = "。".join(f"第{i}条政策内容" for i in range(80)) + "。"
    chunks = ParentChildChunker(parent_size=200, child_size=80, overlap=10).split(
        [ParsedSection("测试政策", content)]
    )
    assert len(chunks) > 1
    assert all(len(parent.content) <= 200 for parent in chunks)
    assert all(len(child) <= 80 for parent in chunks for child in parent.children)
    assert all(parent.heading == "测试政策" for parent in chunks)


def test_internal_document_parse_is_immediately_searchable(client: TestClient) -> None:
    headers = authenticate(client)
    knowledge_base_id = client.post(
        "/api/v1/knowledge-bases",
        headers=headers,
        json={
            "name": "内部知识库",
            "description": "内部办税资料",
            "kb_type": "internal",
        },
    ).json()["data"]["id"]
    document_id = upload_document(
        client,
        headers,
        knowledge_base_id,
        "manual.md",
        "# 申报流程\n第一步登录系统。\n第二步填写申报表。".encode(),
        "text/markdown",
    )

    response = client.post(
        f"/api/v1/documents/{document_id}/parse", headers=headers, json={}
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["parse_status"] == "completed"
    assert data["searchable"] is True
    assert data["parent_chunk_count"] == 1


def test_policy_document_requires_complete_metadata(client: TestClient) -> None:
    headers = authenticate(client)
    knowledge_base_id = create_knowledge_base(client, headers).json()["data"]["id"]
    document_id = upload_document(
        client,
        headers,
        knowledge_base_id,
        "policy.md",
        "# 增值税优惠\n符合条件的小微企业可以适用优惠。".encode(),
        "text/markdown",
    )
    parsed = client.post(f"/api/v1/documents/{document_id}/parse", headers=headers, json={})
    assert parsed.json()["data"]["searchable"] is False

    metadata = client.put(
        f"/api/v1/documents/{document_id}/policy-metadata",
        headers=headers,
        json={
            "policy_title": "增值税优惠政策",
            "doc_no": "财税〔2026〕1号",
            "region": "全国",
            "tax_type": "增值税",
            "taxpayer_type": "小规模纳税人",
            "publish_date": "2026-01-01",
            "effective_start": "2026-01-01",
            "effective_end": "2026-12-31",
            "policy_status": "active",
            "source_url": "https://www.chinatax.gov.cn/example",
        },
    )
    assert metadata.status_code == 200
    assert metadata.json()["data"]["is_complete"] is True

    chunks = client.get(f"/api/v1/documents/{document_id}/chunks", headers=headers)
    assert chunks.status_code == 200
    assert chunks.json()["data"][0]["children"][0]["vector_status"] == "pending"


def test_invalid_policy_period_is_rejected(client: TestClient) -> None:
    headers = authenticate(client)
    knowledge_base_id = create_knowledge_base(client, headers).json()["data"]["id"]
    document_id = upload_document(
        client, headers, knowledge_base_id, "policy.txt", b"policy text", "text/plain"
    )
    client.post(f"/api/v1/documents/{document_id}/parse", headers=headers, json={})

    response = client.put(
        f"/api/v1/documents/{document_id}/policy-metadata",
        headers=headers,
        json={"effective_start": "2026-12-31", "effective_end": "2026-01-01"},
    )
    assert response.status_code == 400
    assert response.json()["code"] == "INVALID_EFFECTIVE_PERIOD"


def test_legacy_doc_parse_failure_is_recorded(client: TestClient) -> None:
    headers = authenticate(client)
    knowledge_base_id = create_knowledge_base(client, headers).json()["data"]["id"]
    document_id = upload_document(
        client, headers, knowledge_base_id, "legacy.doc", b"legacy", "application/msword"
    )
    response = client.post(f"/api/v1/documents/{document_id}/parse", headers=headers, json={})
    assert response.status_code == 400
    assert response.json()["code"] == "DOCUMENT_PARSE_FAILED"
